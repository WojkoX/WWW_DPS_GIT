from asyncio import run
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from datetime import date
import tempfile

from services.nutrition_stats_service import get_nutrition_stats_grouped_by_floor


# =========================================================
# GŁÓWNA FUNKCJA GENERUJĄCA DOCX DLA JEDNEGO PIĘTRA
# =========================================================
def build_floor_docx(floor):

    stats = get_nutrition_stats_grouped_by_floor()

    if floor not in stats:
        raise ValueError(f"Brak danych dla piętra: {floor}")

    data = stats[floor]
    doc = Document()

    section = doc.sections[0]

    # minimalne, ale bezpieczne do druku
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # =====================================================
    # NAGŁÓWEK
    # =====================================================
    header = doc.add_paragraph()
    run = header.add_run(
        f"PIĘTRO: {floor}\t\t\t\t\t\t\t\t\t\tDATA: {date.today().isoformat()}"
    )
    run.bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph("Kierownik Sekcji Żywienia")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    title = doc.add_paragraph(
        "INFORMACJA O DIETACH MIESZKAŃCÓW DPS KOMBATANT\n"
        "SPOŻYWAJĄCYCH POSIŁKI W POKOJACH"
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    # =====================================================
    # TABELA OBECNOŚCI – UKŁAD POZIOMY
    # =====================================================
    presence = data["presence_summary"]

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"

    _presence_cell_text(table.cell(0, 0), "Liczba mieszkańców ogółem")
    _presence_cell_text(table.cell(0, 1), "Osoby będące w szpitalu")
    _presence_cell_text(table.cell(0, 2), "Osoby będące na przepustce")

    _presence_cell_text(table.cell(1, 0), presence["total"], bold=True)
    _presence_cell_text(table.cell(1, 1), presence["hospital"], bold=True)
    _presence_cell_text(table.cell(1, 2), presence["pass"], bold=True)

    # =====================================================
    # WYKAZ DIET
    # =====================================================
    p = doc.add_paragraph("\nWYKAZ DIET:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    hdr[0].text = "Lp."
    hdr[1].text = "Nazwisko i Imię Mieszkańca"
    hdr[2].text = "Nr Pok."
    hdr[3].text = "Rodzaj Diety"

    for idx, r in enumerate(data["present_residents"], start=1):
        row = t.add_row().cells

        row[0].text = str(idx)
        row[1].text = r["name"]
        row[2].text = str(r["room"])

        # ==============================
        # RODZAJ DIETY – FORMAT DWULINIOWY
        # ==============================
        cell = row[3]
        p = cell.paragraphs[0]
        p.clear()

        # linia 1: nazwa + posiłki
        line1 = r["diet"]
        if r["meals"]:
            line1 += f" ({', '.join(r['meals'])})"

        p.add_run(line1)

        # linia 2: UWAGI (BOLD)
        if r["notes"]:
            p.add_run("\n")
            run_notes = p.add_run(r["notes"])
            run_notes.bold = True

    # --- szerokości kolumn ---
    t.autofit = False
    col_widths = [
        Cm(1.2),
        Cm(6.0),
        Cm(1.8),
        Cm(10.3),
    ]
    for row in t.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    # =====================================================
    # ZESTAWIENIA
    # =====================================================
    doc.add_page_break()
    doc.add_paragraph(
        "\nZESTAWIENIE WSZYSTKICH MIESZKAŃCÓW SPOŻYWAJĄCYCH POSIŁKI W POKOJACH:\n"
    ).runs[0].bold = True
    _add_summary_table(doc, data["summary_eat_in_room"])

    doc.add_paragraph(
        "\nZESTAWIENIE W PRZYPADKU JAK WSZYSCY MIESZKAŃCY SPOŻYWAJĄ POSIŁKI W POKOJACH:\n"
    ).runs[0].bold = True
    _add_summary_table(doc, data["summary_all_present"])

    # =====================================================
    # UWAGI
    # =====================================================
    doc.add_paragraph("\n")
    uwagi_table = doc.add_table(rows=1, cols=1)
    uwagi_table.style = "Table Grid"
    cell = uwagi_table.cell(0, 0)

    p = cell.paragraphs[0]
    run = p.add_run("UWAGI:")
    run.bold = True

    for _ in range(8):
        cell.add_paragraph("")

    # =====================================================
    # DATA / SPORZĄDZIŁA
    # =====================================================
    doc.add_paragraph("\n")
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.autofit = True

    sign_table.cell(0, 0).paragraphs[0].add_run("Data").bold = True
    sign_table.cell(0, 1).paragraphs[0].add_run("Sporządził/a").bold = True

    sign_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =====================================================
    # ZAPIS
    # =====================================================
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# =========================================================
# FUNKCJE POMOCNICZE
# =========================================================

def _summary_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold


def _add_summary_table(doc, summary_data):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    _summary_cell(table.cell(0, 0), "Rodzaj diety", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    for col, text in enumerate(
        ["Śniadanie", "Obiad", "Kolacja", "w tym bezmleczna"], start=1
    ):
        _summary_cell(table.cell(0, col), text, bold=True)

    total_b = total_l = total_d = total_m = 0

    for diet, s in summary_data.items():
        row = table.add_row().cells
        _summary_cell(row[0], diet.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _summary_cell(row[1], s["śniadanie"])
        _summary_cell(row[2], s["obiad"])
        _summary_cell(row[3], s["kolacja"])
        _summary_cell(row[4], s["bezmleczna"])

        total_b += s["śniadanie"]
        total_l += s["obiad"]
        total_d += s["kolacja"]
        total_m += s["bezmleczna"]

    row = table.add_row().cells
    _summary_cell(row[0], "RAZEM", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _summary_cell(row[1], total_b, bold=True)
    _summary_cell(row[2], total_l, bold=True)
    _summary_cell(row[3], total_d, bold=True)
    _summary_cell(row[4], total_m, bold=True)


def _presence_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
